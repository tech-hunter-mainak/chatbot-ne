from pathlib import Path
from datetime import datetime

from docx import Document

from .baseParser import BaseParser


class DocxParser(BaseParser):

    def parse(self, filePath: Path) -> dict:

        document = Document(filePath)

        textParts = []

        title = ""

        # Read paragraphs
        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if not text:
                continue

            if not title:
                title = text

            textParts.append(text)

        # Read tables
        for table in document.tables:

            rows = []

            for row in table.rows:

                cells = []

                for cell in row.cells:
                    cells.append(cell.text.strip())

                rows.append(" | ".join(cells))

            if rows:
                textParts.append("\n".join(rows))

        if not title:
            title = filePath.stem.replace("_", " ").replace("-", " ")

        properties = document.core_properties

        stat = filePath.stat()

        return {

            "title": title,

            "text": "\n\n".join(textParts),

            "source": str(filePath),

            "file_name": filePath.name,

            "file_type": "docx",

            "extension": ".docx",

            "author": properties.author or "",

            "created": datetime.fromtimestamp(
                stat.st_ctime
            ).isoformat(),

            "modified": datetime.fromtimestamp(
                stat.st_mtime
            ).isoformat(),

            "size": stat.st_size
        }